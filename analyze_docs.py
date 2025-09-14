#!/usr/bin/env python3
"""
Analyze a directory tree of evidence/legal docs and emit:
- master_review.csv (metadata + smart tags)
- duplicates.csv (hash groups)
Safety: Read-only. No file moves/deletes.
"""
from __future__ import annotations
import csv
import hashlib
import os
import re
import sys
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Tuple

# Optional deps
try:
    import PyPDF2  # type: ignore
except Exception:
    PyPDF2 = None
try:
    import docx  # python-docx
except Exception:
    docx = None

SUPPORTED_EXTS = {
    ".pdf",
    ".doc",
    ".docx",
    ".txt",
    ".rtf",
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
    ".xlsx",
    ".csv",
}

# Heuristic category rules (filename-only; adjust freely)
CATEGORY_RULES = [
    (r"\bCPS\b.*\bComplaint\b|\bComplaint\b.*\bCPS\b", "CPS Complaint"),
    (r"CPS\s*Investigation|\bInvestigation Report\b", "CPS Investigation"),
    (
        r"YWCA|Forensic|Nurse|Spectrum|Helen\s*DeVos|Aversys|Lab\s*Report|DNA|BRAINS|CMH|Oetman|MyChart|Hospital|Pediatric",
        "Medical / Forensic",
    ),
    (r"\bMSP\b|\bCSD\b|Police|Hostage|Assault|Report\b", "Police Report / Incident"),
    (
        r"Order|Recommendation|Referee|Judge|Hearing|Support|Motion|Petition|FOC|GAL|Court",
        "Court / Legal",
    ),
    (
        r"Instructions|Checklist|Notes|Cover\s*Letter|Packet|Summary|Index|Tracker|Guide|Master\s*Plan|Process",
        "External / Instructional",
    ),
]

CHILD_TAGS = [
    "Jace",
    "Josh",
    "Joshua",
    "Nicholas",
    "Nick",
    "Peyton",
    "Owen",
    "Eleanora",
    "Lucas",
    "Jon",
    "John",
]

# Flags for extreme importance (filename heuristics)
EXTREME_PATTERNS = [
    (r"Aversys|DNA|\bat least two donors\b", "DNA / Lab evidence"),
    (r"Noel|Gladding", "Noel-related misconduct"),
    (r"YWCA|Forensic|Nurse", "Forensic nurse / YWCA"),
    (r"Hostage\s*2018|2018.*Hostage|\bHostage\b", "2018 Hostage incident"),
    (r"CMH|PTSD|Mental\s*Health", "CMH / PTSD"),
    (
        r"CPS\s*Complaint\s*(2018|2019|2020)|\b2018\b|\b2019\b|\b2020\b.*CPS\s*Complaint",
        "2018-2020 CPS complaint series",
    ),
]

PROCEDURAL_HINTS = r"Notice of Hearing|Certificate of Service|Proof of Service|Scheduling|Adjourn|Clerk|Register of Actions"


@dataclass
class FileRow:
    rel_path: str
    name: str
    ext: str
    size_bytes: int
    modified: str
    created: str
    sha256: str
    file_type: str  # pdf/doc/image/etc
    pages_or_len: str  # page count for pdf/docx, else ''
    category_guess: str
    child_tags: str
    extreme_flags: str
    suggested_inclusion: str  # YES/NO/TBD
    notes: str
    ocr_needed: str  # YES/NO/blank
    text_preview: str  # first ~500 chars if available


def sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def file_type_from_ext(ext: str) -> str:
    e = ext.lower()
    if e == ".pdf":
        return "pdf"
    if e in {".doc", ".docx"}:
        return "doc"
    if e in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        return "image"
    if e in {".xlsx", ".csv"}:
        return "sheet"
    return "text" if e in {".txt", ".rtf"} else e.lstrip(".")


def count_pdf_pages(p: Path) -> str:
    if not PyPDF2:
        return ""
    try:
        with p.open("rb") as f:
            reader = PyPDF2.PdfReader(f)
            return str(len(reader.pages))
    except Exception:
        return ""


def docx_len(p: Path) -> str:
    if not docx or p.suffix.lower() != ".docx":
        return ""
    try:
        d = docx.Document(str(p))
        # Approximate: paragraphs count
        return str(max(1, len(d.paragraphs)))
    except Exception:
        return ""


def extract_pdf_preview(p: Path, max_chars: int = 500) -> str:
    if not PyPDF2:
        return ""
    try:
        with p.open("rb") as f:
            reader = PyPDF2.PdfReader(f)
            if not reader.pages:
                return ""
            try:
                txt = reader.pages[0].extract_text() or ""
            except Exception:
                txt = ""
            txt = re.sub(r"\s+", " ", txt).strip()
            return txt[:max_chars]
    except Exception:
        return ""


def extract_docx_preview(p: Path, max_chars: int = 500) -> str:
    if not docx or p.suffix.lower() != ".docx":
        return ""
    try:
        d = docx.Document(str(p))
        parts: List[str] = []
        for para in d.paragraphs:
            if para.text:
                parts.append(para.text)
            if sum(len(s) for s in parts) >= max_chars:
                break
        txt = " ".join(parts)
        txt = re.sub(r"\s+", " ", txt).strip()
        return txt[:max_chars]
    except Exception:
        return ""


def extract_text_head(p: Path, max_chars: int = 500) -> str:
    """Best-effort text read for .txt/.rtf and others as plain text."""
    try:
        # Try utf-8 then latin-1 fallback
        try:
            data = p.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            data = p.read_text(encoding="latin-1", errors="ignore")
        data = re.sub(r"\s+", " ", data).strip()
        return data[:max_chars]
    except Exception:
        return ""


def guess_category(name: str) -> str:
    for pat, label in CATEGORY_RULES:
        if re.search(pat, name, re.IGNORECASE):
            return label
    return "Uncategorized"


def tag_children(name: str) -> List[str]:
    found = []
    for child in CHILD_TAGS:
        if re.search(rf"\b{re.escape(child)}\b", name, re.IGNORECASE):
            found.append(child)
    return found


def extreme_flags(name: str) -> List[str]:
    flags = []
    for pat, label in EXTREME_PATTERNS:
        if re.search(pat, name, re.IGNORECASE):
            flags.append(label)
    return flags


def is_procedural(name: str) -> bool:
    return re.search(PROCEDURAL_HINTS, name, re.IGNORECASE) is not None


def walk_files(root: Path) -> List[Path]:
    allp: List[Path] = []
    for dirpath, _, filenames in os.walk(root):
        for fn in filenames:
            p = Path(dirpath) / fn
            if p.suffix.lower() in SUPPORTED_EXTS:
                allp.append(p)
    return allp


def analyze(root: Path) -> Tuple[List[FileRow], Dict[str, List[FileRow]]]:
    rows: List[FileRow] = []
    dup_map: Dict[str, List[FileRow]] = {}
    for p in walk_files(root):
        try:
            stat = p.stat()
            name = p.name
            rel = str(p.relative_to(root))
            ext = p.suffix
            size = stat.st_size
            modified = datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds")
            created = datetime.fromtimestamp(stat.st_ctime).isoformat(timespec="seconds")
            digest = sha256_file(p)
            ftype = file_type_from_ext(ext)
            pages_len = ""
            if ftype == "pdf":
                pages_len = count_pdf_pages(p)
            elif ftype == "doc":
                pages_len = docx_len(p)

            # Text preview + OCR-needed heuristic
            preview = ""
            ocr_needed = ""
            if ftype == "pdf":
                preview = extract_pdf_preview(p)
                try:
                    pg = int(pages_len) if pages_len else 0
                except Exception:
                    pg = 0
                if pg > 0 and len(preview) < 5:
                    ocr_needed = "YES"
                elif len(preview) >= 5:
                    ocr_needed = "NO"
            elif ftype == "doc":
                if p.suffix.lower() == ".docx":
                    preview = extract_docx_preview(p)
                else:
                    preview = ""  # legacy .doc unsupported for preview without extra libs
                if preview:
                    ocr_needed = "NO"
            elif ftype == "text":
                preview = extract_text_head(p)
                if preview:
                    ocr_needed = "NO"
            category = guess_category(name)
            kids = ", ".join(tag_children(name))
            flags = ", ".join(extreme_flags(name))
            # Inclusion heuristic
            if is_procedural(name):
                inclusion = "NO - Procedural only"
            else:
                inclusion = "TBD"

            row = FileRow(
                rel_path=rel,
                name=name,
                ext=ext,
                size_bytes=size,
                modified=modified,
                created=created,
                sha256=digest,
                file_type=ftype,
                pages_or_len=pages_len,
                category_guess=category,
                child_tags=kids,
                extreme_flags=flags,
                suggested_inclusion=inclusion,
                notes="",
                ocr_needed=ocr_needed,
                text_preview=preview,
            )
            rows.append(row)
            dup_map.setdefault(digest, []).append(row)
        except Exception as e:
            sys.stderr.write(f"[warn] failed {p}: {e}\n")
    return rows, dup_map


def write_csv(path: Path, rows: List[Dict[str, str]]):
    if not rows:
        return
    keys = list(rows[0].keys())
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=keys)
        w.writeheader()
        for r in rows:
            w.writerow(r)


def main():
    root = Path(sys.argv[1]) if len(sys.argv) > 1 else Path.cwd()
    out_master = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("master_review.csv")
    out_dups = Path(sys.argv[3]) if len(sys.argv) > 3 else Path("duplicates.csv")
    print(f"[info] Scanning: {root}")
    rows, dup_map = analyze(root)

    # Mark duplicates: keep largest per hash, others suggest NO - Duplicate
    for digest, group in dup_map.items():
        if len(group) > 1:
            largest = max(group, key=lambda r: r.size_bytes)
            for r in group:
                if r is not largest and "NO - Procedural only" not in r.suggested_inclusion:
                    r.suggested_inclusion = "NO - Duplicate"

    # Emit master
    write_csv(out_master, [asdict(r) for r in rows])

    # Emit duplicates (expanded)
    dup_rows: List[Dict[str, str]] = []
    for digest, group in dup_map.items():
        if len(group) > 1:
            for r in group:
                dup_rows.append({
                    "sha256": digest,
                    "name": r.name,
                    "rel_path": r.rel_path,
                    "size_bytes": str(r.size_bytes),
                    "category_guess": r.category_guess,
                    "extreme_flags": r.extreme_flags,
                })
    write_csv(out_dups, dup_rows)
    print(f"[done] Wrote {out_master} ({len(rows)} files) and {out_dups}")


if __name__ == "__main__":
    main()
